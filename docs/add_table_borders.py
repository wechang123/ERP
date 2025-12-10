from docx import Document
from docx.table import Table
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

def add_table_borders(doc_path, output_path):
    doc = Document(doc_path)

    for table in doc.tables:
        # 테이블에 테두리 추가
        tbl = table._tbl
        tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(r'<w:tblPr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')

        # 테두리 스타일 정의
        tblBorders = parse_xml(
            r'<w:tblBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
            r'<w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            r'<w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            r'<w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            r'<w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            r'<w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            r'<w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            r'</w:tblBorders>'
        )

        # 기존 테두리 제거하고 새로 추가
        existing_borders = tblPr.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tblBorders')
        if existing_borders is not None:
            tblPr.remove(existing_borders)
        tblPr.append(tblBorders)

        if tbl.tblPr is None:
            tbl.insert(0, tblPr)

    doc.save(output_path)
    print(f"저장 완료: {output_path}")
    print(f"총 {len(doc.tables)}개 테이블에 테두리 추가됨")

if __name__ == "__main__":
    add_table_borders(
        "ERP_프로젝트_개발보고서.docx",
        "ERP_프로젝트_개발보고서_테두리.docx"
    )
